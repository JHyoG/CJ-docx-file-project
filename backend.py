import docx
import pdb
import numpy as np
import difflib
import re
from docx import Document
from termcolor import colored
from docx.shared import RGBColor
import copy
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.table import _Cell
from collections import Counter
from docx.enum.text import WD_ALIGN_PARAGRAPH

def preprocess(arr):
    if "" in arr:
        arr.remove("")
    if ' ' in arr:
        arr.remove(' ')
    if '' in arr:
        arr.remove('')
    return arr


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    ret = '\n'.join(fullText)
    ret = ret.split('\n')
    return ret


def get_index(input_string, search_string_list, replace_string):
    result = ""
    index = 0
    index_arr = []

    prev = 0
    while index < len(input_string):
        for word in search_string_list:
            next_index = input_string.find(word, index)

            if next_index == -1:
                result += input_string[index:]
                break

            result += input_string[index:next_index]

            result += replace_string

            index = next_index + len(word)

            index_arr.append([prev, next_index, len(word)])
            prev = index
        break
    return result, index_arr




def what_first(text):

    for i in text:
        if i == '+':
            return True
        elif i == '-':
            return False

def color_index(text):
    minus_indices = [i for i, char in enumerate(text) if char == '-']
    plus_indices = [i for i, char in enumerate(text) if char == '+']

    return minus_indices, plus_indices

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

# 입력 -> 원본파일 테이블, (정형화된)만든 테이블 , 출력 -> 색이 적용된 테이블
def extract_cell_colors(source_table,target_table):

    # 정규 표현식 패턴 컴파일
    pattern = re.compile('w:fill=\"(\\S*)\"')


    # 색상 정보 저장할 딕셔너리
    color_map = {}

    # 테이블 순회
    for row_index, row in enumerate(source_table.rows):
        for col_index, cell in enumerate(row.cells):
                # 셀의 XML 텍스트 가져오기
                cell_xml = cell._element.xml

                # 정규 표현식을 사용하여 w:fill 속성에서 색상 추출
                match = pattern.search(cell_xml)
                if match:
                    color = match.group(1)

                    # 색상 정보 저장
                    color_map[(row_index, col_index)] = color



    # target_table에 색상 적용

    for row_index, row in enumerate(target_table.rows):
        for col_index, cell in enumerate(row.cells):
            # 행과 열의 인덱스로 딕셔너리에서 색상을 가져옴
            if (row_index, col_index) in color_map:

                color = color_map[(row_index, col_index)]

                # 현재 셀의 XML 엘리먼트를 가져옴
                cell_xml_element = cell._tc
                # 셀의 특성 엘리먼트를 가져오거나 생성
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                # 배경색을 나타내는 XML 엘리먼트 생성
                shape_obj = OxmlElement('w:shd')
                shape_obj.set(qn('w:fill'), color)
                # 배경색을 셀의 특성에 추가
                table_cell_properties.append(shape_obj)

    return target_table

# 입력 -> 테이블객체, 출력 -> 테이블 객체(원본 format돼있는 형식) , 방문하지 않을 셀 정보(행,열 정보)-> ex)[(2,1),(4,1)]
def get_table_format(table_object):
    out_doc = Document()  # 새로운 Word 문서 생성
    paragraph = out_doc.add_paragraph()
    doc1_table_copy = copy.deepcopy(table_object)

    table_cell_info=np.array([])

    # 원본 테이블의 정보(ex)병합된 셀의 정보)를 저장
    for i, row in enumerate(table_object.rows): # 행 인덱스, 행정보 저장
        for j, cell in enumerate(row.cells): # 행의 모든 셀에 대해 반복, cell에는 현제 셀에대한 정보저장, j에는 해당셀의 열 인덱스 저장
            c=table_object.cell(i,j)
            cell_info = np.array([str(c._tc.top), str(c._tc.bottom),str(c._tc.left),str(c._tc.right)])
            result=''.join(cell_info)
            table_cell_info = np.append(table_cell_info, [result])

    count_dict = Counter(table_cell_info)


    duplicates = [item for item, count in count_dict.items() if count > 1]
    # 원본형태의 행,열 개수를 가진 '정형화된' 테이블 만들기
    num_rows = len(doc1_table_copy.rows) # 행
    num_cols = len(doc1_table_copy.columns) # 열

    table_cell_info = np.reshape(table_cell_info,(num_rows, num_cols))
    table=out_doc.add_table(rows=num_rows,cols=num_cols)
    table.style = "Table Grid"
    table= extract_cell_colors(table_object,table)

    duplicate_cell_positions = {}

    for i in range(num_rows):
        for j in range(num_cols):
            value_to_check = table_cell_info[i, j]

            if value_to_check in duplicates:
                if value_to_check not in duplicate_cell_positions:
                    duplicate_cell_positions[value_to_check] = []
                duplicate_cell_positions[value_to_check].append((i, j))

    not_visit_cell_info = [] # 방문하지 않을 셀 리스트(중복된 값을 넣지 않기 위함)

    for key, positions in duplicate_cell_positions.items():
        rest_cells = positions[1:]
        not_visit_cell_info.extend(rest_cells)

    #  중복된 값에 대해 최소값과 최대값을 가지는 위치 추출

    min_max_positions = {}

    for value, positions in duplicate_cell_positions.items():
        min_row_col = positions[0] # => min(positions, key=lambda x: x[0] + x[1])
        max_row_col = positions[-1] # => max(positions, key=lambda x: x[0] + x[1])
        min_max_positions[value] = {'min': min_row_col, 'max': max_row_col}


    # merge 해주기

    for value, positions in min_max_positions.items():
        min_value = positions['min']
        max_value = positions['max']

        if min_value[0] == max_value[0] and min_value[1] != max_value[1]: # horizental merged일때
            colA=table.columns[min_value[1]]
            colB=table.columns[max_value[1]]
            colA.cells[min_value[0]].merge(colB.cells[min_value[0]])
            # print(f"값 '{value}'의 수평 병합: 행 {min_value[0]}, 열 {min_value[1]} ~ 행 {max_value[0]}, 열 {max_value[1]}")

        elif min_value[0] != max_value[0] and min_value[1] == max_value[1]:# vertical merged일때
            RowA=table.rows[min_value[0]]
            RowB=table.rows[max_value[0]]
            RowA.cells[min_value[1]].merge(RowB.cells[min_value[1]])
            # print(f"값 '{value}'의 수직 병합: 행 {min_value[0]}, 열 {min_value[1]} ~ 행 {max_value[0]}, 열 {max_value[1]}")

        else: # horizental + vertical merged일때
            RowA=table.rows[min_value[0]]
            RowB=table.rows[max_value[0]]
            RowA.cells[min_value[1]].merge(RowB.cells[max_value[1]])

    return table , not_visit_cell_info


def get_table(input_file, compare_file, out_doc, index):
    doc1 = docx.Document(input_file)
    doc2 = docx.Document(compare_file)
    doc1_table = doc1.tables[index]
    doc2_table = doc2.tables[index]
    text1_arr = []
    text2_arr = []
    differences = []
    paragraph = out_doc.add_paragraph()

    # formatted_table은 테이블 객체 ,not_visit_cell_arr은 리스트
    formatted_table, not_visit_cell_arr= get_table_format(doc2_table)
    d=difflib.Differ()
    for i, (table1_row, table2_row) in enumerate(zip(doc1_table.rows, doc2_table.rows)):
        text1 = [cell1.text for cell1 in table1_row.cells]
        text2 = [cell2.text for cell2 in table2_row.cells]
        text1_arr.append(text1)
        text2_arr.append(text2)
        differences.append([''.join(d.compare(str1, str2)).replace('  ','') for str1, str2 in zip(text1, text2)])

    da_table = formatted_table._tbl
    paragraph._p.addnext(da_table)

    red = RGBColor(255, 0, 0)
    green = RGBColor(82, 167, 54)


    for row in range(len(doc1_table.rows)):
        for col in range(len(doc1_table.columns)):
            current_cell = (row, col) # not_visit_cell_arr안에 있는 행,열은 넘어가기
            if current_cell in not_visit_cell_arr:
                continue

            cell = formatted_table.cell(row, col)

            if  text1_arr[row][col] == text2_arr[row][col]:  #  원본 테이블과 바뀐게 없으면
                cell.text = text2_arr[row][col]
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

            else:
                paragraph = cell.paragraphs[0]
                text = differences[row][col].rstrip().lstrip()

                for count in range(len(text)):
                    if count >= 2 and text[count-2] in ('-', '+') and text[count-1] == ' ':
                        colored_run = paragraph.add_run(text[count])
                        colored_run.font.color.rgb = red if text[count-2] == '-' else green
                        colored_run.font.strike = True
                    elif text[count] in ('+', '-') and ord(text[count+1]) != 32:
                        paragraph.add_run(text[count])
                    elif text[count] not in ('+', '-') and (count == 0 or text[count-1] not in ('+', '-')):
                        paragraph.add_run(text[count])

                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    return paragraph

def create_document_data(input_string,compare_string,org_text, edit_text,output_name, table_index,table_count):

    output_doc = Document()
    style = output_doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    differences = []
    text_count=0
    tbl_count=0

    d=difflib.Differ()

    preprocess(input_string)
    preprocess(compare_string)

    #pdb.set_trace()
    for i, (input_text, compare_text) in enumerate(zip(input_string, compare_string)):
        differences.append([''.join(d.compare(input_text, compare_text)).replace('  ','') ])

    red = RGBColor(255, 0, 0)
    green = RGBColor(82, 167, 54)

    for i in range(len(compare_string) + table_count):
        if i in table_index:
            paragraph = get_table(org_text, edit_text, output_doc, tbl_count)
            tbl_count+=1
        elif input_string[text_count]==compare_string[text_count]:
            paragraph = output_doc.add_paragraph(compare_string[text_count])
            text_count+=1
        else:
            paragraph = output_doc.add_paragraph()
            text = differences[text_count][0].rstrip().lstrip()
            for count in range(len(text)):
                if count >= 2 and text[count-2] in ('-', '+') and text[count-1] == ' ':
                    colored_run = paragraph.add_run(text[count])
                    colored_run.font.color.rgb = red if text[count-2] == '-' else green
                    if colored_run.font.color.rgb == red:
                        colored_run.font.strike = True
                elif text[count] in ('+', '-') and ord(text[count+1]) != 32:
                    paragraph.add_run(text[count])
                elif text[count] not in ('+', '-') and (count == 0 or text[count-1] not in ('+', '-')):
                    paragraph.add_run(text[count])
            text_count+=1
    output_doc.save(output_name) 
    return output_name


def compareList(l1, l2):
    return [i==j for i, j in zip(l1, l2)]

def get_table_index(src):
    table_index = []
    count = 0
    doc1 = docx.Document(src)

    for child in doc1.element.body.xpath('w:p | w:tbl'):
        print(f'child : {child}')
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc1)
            if paragraph.text=='':
                continue
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc1)
            table_index.append(count)
        count+=1

    return table_index

# 전체 실행
def start(edit_text, org_text,path):
    path = path +'/'
    original = getText(org_text)
    edit_getText = getText(edit_text) ## 전처리
    name = org_text.split('\\')[-1].split('.')[0] 
    table_index = get_table_index(org_text)  
    table_count=len(table_index)
    print("##수정 test cast 검사중##")
    output = create_document_data(original,edit_getText,org_text,edit_text,path+name+'_result.docx',table_index,table_count)
    print("##검사완료###")
    return output



    